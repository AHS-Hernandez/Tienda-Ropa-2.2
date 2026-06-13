# -*- coding: utf-8 -*-
"""
Genera la documentacion formal del Modulo NoSQL (MongoDB) en formato .docx.
Sin colores fuertes; verde pastel sutil para titulos. Estilo formal y defendido.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VERDE_TITULO = RGBColor(0x4A, 0x6B, 0x4F)   # verde pastel oscuro
VERDE_SUBTI  = RGBColor(0x6B, 0x8E, 0x70)
GRIS_TXT     = RGBColor(0x33, 0x33, 0x33)
GRIS_TABLA   = "E8EEE9"                      # gris-verdoso muy claro

doc = Document()

# --- margenes y fuente base ---
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = GRIS_TXT

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = VERDE_TITULO
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(8)

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = VERDE_TITULO
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = VERDE_SUBTI
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

def parrafo(text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    return p

def vinieta(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p

def code_block(code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    for i, line in enumerate(code_lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x22, 0x33, 0x22)
    return p

def tabla(headers, filas):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = VERDE_TITULO
        set_cell_bg(hdr[i], GRIS_TABLA)
    for fila in filas:
        row = t.add_row().cells
        for i, val in enumerate(fila):
            row[i].text = ""
            para = row[i].paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = GRIS_TXT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# =====================================================================
# PORTADA
# =====================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n\nTiendaRopa\n")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = VERDE_TITULO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Modulo NoSQL — MongoDB\n")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = VERDE_SUBTI

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documentacion tecnica del componente documental\n")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = GRIS_TXT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n\n\nProyecto integrador — Base de Datos II\n")
r.font.size = Pt(11)
r.font.color.rgb = GRIS_TXT
r2 = p.add_run("Universidad — 2026\n")
r2.font.size = Pt(11)
r2.font.color.rgb = GRIS_TXT

doc.add_page_break()

# =====================================================================
# 1. INTRODUCCION
# =====================================================================
h1("1. Introduccion al sistema y rol del componente NoSQL")

parrafo(
"TiendaRopa es un sistema ERP distribuido para una cadena de tiendas de ropa "
"con presencia en multiples sedes. La plataforma esta construida como aplicacion "
"web en Next.js 15 sobre React 19, con SQL Server como sistema de registro "
"transaccional y MongoDB como componente documental de enriquecimiento. "
"El sistema atiende a cuatro perfiles diferenciados —Dueno, Administrador de Sede, "
"Vendedor y Cliente final— y resuelve la operacion completa del negocio: "
"catalogo maestro, inventario por sede, ventas en linea y presenciales, "
"campanas de marketing, y la experiencia del comprador en la tienda online."
)

parrafo(
"El componente NoSQL no es un reemplazo del motor relacional. Es una pieza "
"deliberada cuya unica responsabilidad es almacenar y procesar la informacion "
"que no encaja en el modelo relacional y que justifica la inversion en un "
"segundo motor. SQL Server sigue siendo el sistema de registro: clientes, "
"ventas, stock, kardex y facturacion residen alli. MongoDB asume las "
"responsabilidades complementarias que requieren esquemas flexibles, "
"escritura masiva y caducidad automatica."
)

h2("1.1 Por que dos motores y no uno")

parrafo(
"La decision de incorporar MongoDB se tomo despues de identificar tres "
"problemas concretos del dominio que el modelo relacional resuelve mal o "
"a un costo desproporcionado:"
)

vinieta(
"Resenas de productos con respuestas embebidas, fotos opcionales y campos "
"de moderacion: una entidad cuyas partes opcionales son la mayoria, donde "
"un modelo relacional generaria tablas casi vacias y JOINs costosos para "
"servir la ficha de producto."
)
vinieta(
"Registro de comportamiento del cliente (vistas, busquedas, agregados al "
"carrito, abandonos): escritura intensa y constante, lectura siempre por "
"ventana de tiempo + dimension, sin necesidad de transaccionalidad estricta."
)
vinieta(
"Feed personalizado por cliente: una lectura O(1) construida a partir de "
"varias senales heterogeneas, que debe caducar y regenerarse sin intervencion "
"manual."
)

parrafo(
"Cada uno de estos problemas tiene una solucion natural en MongoDB: "
"documento auto-contenido con validador, coleccion time-series con TTL, "
"y vista materializada con TTL. Implementarlos en SQL Server seria forzar "
"la herramienta y pagar un costo continuo en mantenimiento y rendimiento."
)

h2("1.2 Estructura del componente NoSQL")

parrafo(
"La base se llama tiendaropa_nosql y contiene tres colecciones productivas, "
"cada una con un patron distinto de Mongo y una razon de negocio distinta:"
)

tabla(
    ["Coleccion", "Patron de Mongo", "Funcion en el negocio"],
    [
        ["resenas", "Documento con validador $jsonSchema, embebidos y referencias",
         "Opiniones de clientes con prueba de compra"],
        ["eventos", "Coleccion time-series con TTL de 90 dias",
         "Log de comportamiento del cliente en la web"],
        ["feed_cliente", "Vista materializada con TTL controlado por vigencia",
         "Recomendaciones personalizadas listas para servir"],
    ]
)

parrafo(
"Los scripts de creacion, carga y consulta estan en la carpeta MongoDB/ "
"del proyecto, numerados del 00 al 06 para garantizar el orden de ejecucion. "
"Las colecciones ref_* que aparecen en los scripts contienen datos de "
"referencia que en produccion provienen de SQL; existen solo para que los "
"scripts puedan ejecutarse y verificarse de forma autocontenida."
)

doc.add_page_break()

# =====================================================================
# 2. COLECCION RESENAS
# =====================================================================
h1("2. Coleccion resenas — Modulo B")

h2("2.1 Justificacion del modelo documental")

parrafo(
"Una resena es una entidad heterogenea por naturaleza. El cliente puede "
"adjuntar fotos o no; el vendedor o el administrador pueden responder "
"o no; el moderador puede ocultar la resena o dejarla publicada. Modelar "
"esto en filas y tablas separadas implica: una tabla de resenas, una de "
"fotos, una de respuestas, y multiples JOINs cada vez que se construye "
"la ficha de producto. El costo es desproporcionado para un caso de uso "
"que se lee con frecuencia y muta poco."
)

parrafo(
"En MongoDB el documento agrupa todo. La resena, sus fotos y sus respuestas "
"viven juntas en un solo objeto. La ficha de producto se sirve con una "
"unica consulta indexada. El campo respuestas es un arreglo de subdocumentos "
"que se extiende sin migrar esquema. Las invariantes que si son obligatorias "
"(rating entre 1 y 5, estado en un enumerado, referencias presentes) se "
"declaran como validador del motor, no como codigo de aplicacion."
)

h2("2.2 Estructura del documento")

parrafo(
"Cada documento de la coleccion resenas contiene las siguientes categorias "
"de campos:"
)

vinieta(
"Referencias al sistema relacional: id_producto, id_cliente, id_venta, id_sede. "
"Son la prueba documental de compra. La aplicacion verifica la combinacion "
"contra Ventas.Venta_Detalle antes de aceptar la insercion."
)
vinieta(
"Datos de la opinion: rating (1 a 5), titulo (hasta 120 caracteres), texto "
"(hasta 2000 caracteres), fotos (arreglo de URLs), votos_util."
)
vinieta(
"Estado y temporalidad: estado (publicada u oculta), fecha de emision."
)
vinieta(
"Respuestas embebidas: arreglo respuestas[] con autor, texto y fecha por "
"cada respuesta del vendedor o moderador."
)

h2("2.3 Validador a nivel de motor")

parrafo(
"La coleccion declara un validador $jsonSchema con validationLevel = strict "
"y validationAction = error. El motor rechaza cualquier insercion que "
"viole las invariantes, sin depender de que la capa de aplicacion se "
"acuerde de validar. Las invariantes declaradas son:"
)

code_block([
"db.createCollection('resenas', {",
"  validator: {",
"    $jsonSchema: {",
"      bsonType: 'object',",
"      required: ['id_producto','id_cliente','id_venta','id_sede',",
"                 'rating','fecha','estado'],",
"      properties: {",
"        rating: { bsonType: 'number', minimum: 1, maximum: 5 },",
"        estado: { enum: ['publicada','oculta'] },",
"        titulo: { bsonType: 'string', maxLength: 120 },",
"        texto:  { bsonType: 'string', maxLength: 2000 }",
"      }",
"    }",
"  },",
"  validationLevel: 'strict',",
"  validationAction: 'error'",
"});",
])

parrafo(
"La declaracion bsonType usa number en lugar de int por una razon practica: "
"los literales numericos en mongosh son double por defecto. Si se exigiera "
"int, el validador rechazaria datos correctos sembrados desde el shell. "
"El driver de Node, que es el que usa la aplicacion productiva, controla "
"los tipos explicitamente y mantiene la coherencia."
)

h2("2.4 Indices y la regla 'una opinion por compra'")

parrafo(
"La coleccion declara tres indices, cada uno asociado a una operacion "
"concreta del negocio:"
)

tabla(
    ["Indice", "Tipo", "Operacion que soporta"],
    [
        ["{id_cliente:1, id_producto:1, id_venta:1}", "Unico",
         "Una resena por compra. Imposible reseñar dos veces la misma venta."],
        ["{id_producto:1, estado:1, fecha:-1}", "Compuesto",
         "Resenas publicadas de un producto, mas recientes primero. Es la lectura mas frecuente."],
        ["{id_producto:1, rating:1}", "Compuesto",
         "Apoya las agregaciones de promedio y distribucion de estrellas."],
    ]
)

parrafo(
"El indice unico es la piedra angular de la regla 'una opinion por compra'. "
"No depende de que la aplicacion verifique nada: aunque dos peticiones "
"concurrentes intenten insertar la misma resena, MongoDB rechaza la segunda. "
"Esta regla de negocio queda garantizada por el motor, no por una capa "
"que pueda olvidarse de validar."
)

h2("2.5 Consultas que usa la aplicacion")

parrafo(
"Las consultas mas representativas estan en el script "
"MongoDB/03_consultas_resenas.js. Las cinco consultas que ejecuta la "
"aplicacion productiva son:"
)

h3("Consulta 1 — Resenas publicadas de un producto")

parrafo(
"Es la consulta mas frecuente del sistema: cada visita a la ficha de "
"producto la dispara. Usa el indice ix_resena_producto, que ordena "
"fisicamente por fecha descendente, eliminando cualquier operador Sort."
)

code_block([
"db.resenas.find(",
"  { id_producto: 1001, estado: 'publicada' },",
"  { _id: 0, id_cliente: 1, rating: 1, titulo: 1, texto: 1, fecha: 1 }",
").sort({ fecha: -1 });",
])

h3("Consulta 2 — Promedio y total por producto")

parrafo(
"Alimenta el campo rating_promedio que el front muestra junto al precio. "
"Agrega por id_producto con $group y $avg. El round(2) garantiza que la "
"presentacion sea estable. Es una agregacion barata gracias al indice "
"ix_resena_rating."
)

code_block([
"db.resenas.aggregate([",
"  { $match: { estado: 'publicada' } },",
"  { $group: { _id: '$id_producto',",
"              rating_promedio: { $avg: '$rating' },",
"              total: { $sum: 1 } }},",
"  { $project: { _id: 0, id_producto: '$_id',",
"                rating_promedio: { $round: ['$rating_promedio', 2] },",
"                total: 1 }},",
"  { $sort: { rating_promedio: -1, total: -1 } }",
"]);",
])

h3("Consulta 3 — Distribucion de estrellas")

parrafo(
"Genera el histograma clasico de reseñas (cuantas de 5 estrellas, cuantas "
"de 4, etc.) que se muestra como barras en la ficha del producto. Demuestra "
"agrupacion por valor de rating."
)

h3("Consulta 4 — Top productos mejor valorados con minimo de reseñas")

parrafo(
"Aplica un filtro de minimo (>=2 resenas) para evitar que un producto con "
"una unica reseña de 5 estrellas domine el ranking. Es un patron clasico "
"de proteccion contra outliers en sistemas de valoracion."
)

h3("Consulta 5 — Resenas ocultas para moderacion")

parrafo(
"Sirve la bandeja del administrador. La diferencia con la consulta 1 es "
"el filtro de estado: el publico solo ve estado='publicada', mientras que "
"el moderador ve estado='oculta'. Una misma coleccion atiende dos vistas "
"de seguridad distintas con solo cambiar el filtro."
)

h2("2.6 Como se integra con SQL Server")

parrafo(
"La aplicacion verifica en SQL que la venta exista, que pertenezca al "
"cliente autenticado y que incluya el producto antes de aceptar la "
"insercion en Mongo. Si la verificacion falla, la insercion se rechaza "
"con el mensaje 'no puedes resenar este producto'. Mongo guarda unicamente "
"las referencias numericas; el nombre del cliente y el del producto se "
"resuelven en consulta vía las funciones getNombresClientes y "
"getNombresProductos. Esta separacion evita que la resena quede "
"desincronizada si el cliente actualiza sus datos."
)

doc.add_page_break()

# =====================================================================
# 3. COLECCION EVENTOS (TIME-SERIES)
# =====================================================================
h1("3. Coleccion eventos — Modulo C, capa 1")

h2("3.1 Por que una coleccion time-series")

parrafo(
"Los eventos de comportamiento del cliente (vista de producto, busqueda, "
"agregado al carrito, abandono de sesion) tienen un patron de acceso "
"caracteristico: la escritura es masiva y continua, y la lectura siempre "
"se hace por ventana de tiempo combinada con una dimension de agrupacion. "
"Las preguntas tipicas son 'productos mas vistos en mi sede esta semana', "
"'busquedas sin resultados en los ultimos 7 dias', 'embudo de conversion "
"por hora'. Nunca se pregunta por un evento individual."
)

parrafo(
"MongoDB time-series esta disenado exactamente para ese patron. El motor "
"agrupa fisicamente los documentos por metaField + timeField y aplica "
"compresion columnar. Las escrituras de log son baratas, las agregaciones "
"temporales son rapidas, y la retencion se controla con TTL automatico."
)

h2("3.2 Configuracion declarada")

code_block([
"db.createCollection('eventos', {",
"  timeseries: {",
"    timeField:   'ts',",
"    metaField:   'meta',    // { id_sede, tipo, id_producto }",
"    granularity: 'seconds'",
"  },",
"  expireAfterSeconds: 7776000  // 90 dias",
"});",
])

parrafo(
"timeField indica el campo temporal sobre el que se ordenan los buckets "
"internos. metaField agrupa las dimensiones que casi siempre acompanan "
"al filtro temporal: sede, tipo de evento, producto. granularity='seconds' "
"se eligio porque las sesiones web generan eventos en el orden de segundos, "
"no de minutos ni de horas; el motor optimiza la compresion en consecuencia."
)

parrafo(
"expireAfterSeconds = 7776000 fija una retencion de 90 dias para el dato "
"crudo. Pasado ese plazo el motor borra los buckets automaticamente. Esta "
"caducidad libera espacio y mantiene la coleccion acotada sin necesidad "
"de un proceso de purga manual. La intencion de negocio (las metricas "
"derivadas y las recomendaciones) sobrevive en feed_cliente y en "
"agregaciones materializadas."
)

h2("3.3 Estructura del documento")

parrafo(
"Las colecciones time-series no admiten validador $jsonSchema, por lo que "
"la forma del documento se documenta y se garantiza en la capa de ingesta. "
"La estructura productiva es:"
)

code_block([
"{",
"  ts: ISODate('2026-06-13T10:32:14Z'),",
"  meta: {",
"    id_sede:     1,",
"    tipo:        'vista_producto',",
"    id_producto: 1001",
"  },",
"  id_cliente: 4,",
"  payload: {",
"    categoria:        'Camisas',",
"    sesion:           'sess_xyz',",
"    termino_busqueda: null,",
"    resultados:       null",
"  }",
"}",
])

parrafo(
"El campo tipo discrimina la naturaleza del evento. Los valores en "
"produccion son: vista_producto, busqueda, agregar_carrito, abandono_carrito "
"y vista_categoria. Cada uno alimenta una metrica distinta."
)

h2("3.4 Indice de apoyo")

parrafo(
"La coleccion time-series indexa automaticamente por meta + tiempo, lo "
"que cubre la mayoria de consultas. Se anadio un indice adicional para "
"el patron 'historial de un producto a lo largo del tiempo':"
)

code_block([
"db.eventos.createIndex(",
"  { 'meta.id_producto': 1, ts: -1 },",
"  { name: 'ix_eventos_producto' }",
");",
])

h2("3.5 Consultas que alimentan el negocio")

parrafo(
"El script MongoDB/04_consultas_eventos.js contiene las consultas "
"productivas. Cada una responde a una pregunta concreta del negocio:"
)

tabla(
    ["Consulta", "Pregunta de negocio"],
    [
        ["Productos mas vistos por sede (7 dias)",
         "Que esta llamando la atencion en cada sucursal"],
        ["Busquedas sin resultados",
         "Que quiere el cliente y no tenemos en catalogo"],
        ["Embudo por tipo de evento (vista -> carrito -> abandono)",
         "Salud del POS web; donde se pierden las conversiones"],
        ["Eventos por hora del dia",
         "Curva de uso para planificar promociones y staffing"],
        ["Vistas por categoria por cliente",
         "Senal de afinidad para alimentar el feed personalizado"],
    ]
)

parrafo(
"Las cinco consultas se ejecutan con $match sobre el campo ts (recibido "
"como ISODate) seguido de $group por la dimension correspondiente. El "
"motor las resuelve con seek directo a los buckets relevantes; no hace "
"falta recorrer la coleccion completa."
)

h2("3.6 Captura de eventos desde la aplicacion")

parrafo(
"La aplicacion Next.js expone el hook useTrackEvent (definido en "
"shared-component-system/hooks/use-track-event.ts) que las paginas del "
"cliente invocan al producirse cada interaccion relevante. El hook envia "
"el evento al endpoint /api/cliente/eventos, que lo inserta en la "
"coleccion. La insercion no es bloqueante para el usuario: el evento "
"viaja en paralelo a la renderizacion."
)

doc.add_page_break()

# =====================================================================
# 4. COLECCION FEED_CLIENTE
# =====================================================================
h1("4. Coleccion feed_cliente — Modulo C, capa 2")

h2("4.1 Patron de vista materializada")

parrafo(
"El feed personalizado es el equivalente en MongoDB de la tabla "
"Inventario.Stock_Actual en SQL: una vista materializada que la "
"aplicacion lee sin computar nada en el momento. Un proceso lee el log "
"de eventos, ejecuta las cuatro senales del algoritmo de recomendacion, "
"y deja el resultado listo en la coleccion. Cuando el cliente entra al "
"sitio, la lectura es instantanea: un find por id_cliente sobre el indice "
"unico devuelve el documento completo."
)

parrafo(
"La alternativa —calcular las recomendaciones en cada request— exigiria "
"recorrer eventos, productos comprados, vistos y stock, en cada visita. "
"Es viable para pocos clientes pero se degrada rapidamente. La "
"materializacion separa el costo computacional del costo de lectura."
)

h2("4.2 Estructura del documento")

parrafo(
"Cada documento agrupa todo lo necesario para servir el feed sin nuevas "
"consultas:"
)

code_block([
"{",
"  id_cliente:       4,",
"  id_sede_contexto: 1,",
"  generado: ISODate('2026-06-13T08:00:00Z'),",
"  vigencia: ISODate('2026-06-13T14:00:00Z'),  // +6 horas",
"  perfil: {",
"    categorias_top: [",
"      { categoria: 'Camisas',    vistas: 12 },",
"      { categoria: 'Pantalones', vistas: 7 }",
"    ]",
"  },",
"  recomendaciones: [",
"    {",
"      id_producto: 1004, nombre: 'Camisa lino azul',",
"      precio: 240, score: 1.0,",
"      fuente: 'vistos_no_comprados',",
"      motivo: 'Porque lo viste y aun no lo compraste'",
"    },",
"    /* ... hasta 12 ... */",
"  ]",
"}",
])

parrafo(
"Los campos generado y vigencia controlan el ciclo de vida. id_sede_contexto "
"indica la sede usada para verificar disponibilidad real del producto: una "
"recomendacion debe poder cumplirse en la sucursal del cliente, no en otra. "
"El perfil resume las categorias mas vistas y se incluye para que el "
"front pueda mostrar 'tus categorias favoritas' sin nuevas consultas."
)

h2("4.3 Indices y TTL")

code_block([
"db.feed_cliente.createIndex(",
"  { id_cliente: 1 }, { unique: true, name: 'ux_feed_cliente' }",
");",
"",
"db.feed_cliente.createIndex(",
"  { vigencia: 1 },",
"  { expireAfterSeconds: 0, name: 'ttl_feed' }",
");",
])

parrafo(
"El indice unico por id_cliente garantiza un feed por cliente: cualquier "
"intento de generar dos en paralelo se resuelve con upsert y termina en "
"un solo documento. El indice TTL con expireAfterSeconds=0 borra el "
"documento en el momento exacto en que vigencia queda en el pasado. La "
"siguiente lectura del cliente activa la regeneracion."
)

parrafo(
"Esta combinacion evita servir recomendaciones obsoletas sin necesidad "
"de un job de limpieza separado. El motor se encarga. La regeneracion "
"queda gobernada por la demanda real: si el cliente no vuelve, no se "
"recalcula nada; si vuelve, se recalcula bajo demanda."
)

h2("4.4 El algoritmo de recomendacion")

parrafo(
"El script MongoDB/05_generar_feed.js implementa cuatro senales con pesos "
"distintos. Cada senal corresponde a una hipotesis distinta sobre la "
"intencion de compra, y el peso refleja cuan directa es esa hipotesis."
)

tabla(
    ["Senal", "Peso", "Hipotesis de negocio"],
    [
        ["vistos_no_comprados", "1.0",
         "Lo miro pero no lo compro; intencion directa de compra pendiente"],
        ["co_visitacion", "0.8",
         "Clientes con comportamiento parecido vieron tambien este producto"],
        ["categoria_afin", "0.6",
         "Afinidad declarada por las propias vistas del cliente"],
        ["trending_sede", "0.4",
         "Contexto local: lo que esta de moda en su sucursal"],
    ]
)

parrafo(
"Las cuatro senales se evaluan por separado, se reunen en un conjunto "
"de candidatos, y luego se dedupe quedandose con el mayor score por "
"producto. Despues se aplican dos reglas duras heredadas de la operacion "
"real:"
)

vinieta(
"No recomendar lo que el cliente ya compro: revisar el feed con productos "
"comprados es ruido, no recomendacion."
)
vinieta(
"No recomendar lo agotado en su sede: la recomendacion debe poder "
"cumplirse. Si el producto no esta disponible localmente, se descarta."
)

parrafo(
"El resultado se ordena por score descendente y se limita a las 12 "
"primeras posiciones. El documento final se materializa con replaceOne + "
"upsert sobre la clave id_cliente. La vigencia se fija en 6 horas, lo "
"que equilibra frescura y costo de regeneracion."
)

h2("4.5 Como sobrevive a la operacion concurrente")

parrafo(
"El algoritmo es idempotente: ejecutar generarFeed(idCliente) dos veces "
"produce el mismo resultado, porque replaceOne reemplaza el documento "
"completo. Esto significa que si dos procesos disparan la regeneracion "
"al mismo tiempo, ganan el ultimo escritor sin corrupcion."
)

parrafo(
"La caducidad del feed esta desacoplada de la operacion del cliente: "
"el TTL se aplica en background, sin tocar el flujo de la peticion web. "
"Si el feed caduca entre la lectura y la renderizacion, el cliente vera "
"un feed regenerado en la siguiente visita, no en la actual; nunca un "
"feed corrupto o a medio armar."
)

doc.add_page_break()

# =====================================================================
# 5. INTEGRACION CON SQL SERVER
# =====================================================================
h1("5. Integracion con SQL Server")

parrafo(
"La separacion de responsabilidades entre los dos motores es estricta. "
"SQL Server posee la verdad transaccional (clientes, productos, ventas, "
"stock); MongoDB posee la verdad documental (resenas, eventos, feed). "
"La integracion se realiza en la capa de aplicacion, no en la base."
)

h2("5.1 Validacion cruzada al insertar")

parrafo(
"Antes de insertar una resena en Mongo, el endpoint /api/cliente/resenas "
"consulta SQL para verificar tres condiciones:"
)

vinieta("El cliente autenticado es el dueno de la venta (id_cliente coincide).")
vinieta("La venta existe y esta en estado 'Pagada' o 'Entregada' (nunca 'Borrador').")
vinieta("La venta incluye el producto que se intenta resenar.")

parrafo(
"Si alguna condicion falla, la insercion se rechaza. El validador "
"$jsonSchema de Mongo es la red de seguridad final, no la primera linea "
"de defensa."
)

h2("5.2 Enriquecimiento al servir")

parrafo(
"Las resenas se almacenan con identificadores numericos (id_cliente, "
"id_producto). Al servirlas al frontend, la aplicacion resuelve los "
"nombres consultando SQL:"
)

vinieta("getNombresClientes(ids) -> Persona.Cliente JOIN Persona.Persona")
vinieta("getNombresProductos(ids) -> Producto.Producto")

parrafo(
"Esta resolucion en consulta evita que la resena quede desincronizada si "
"el cliente actualiza sus datos o si el nombre del producto cambia. El "
"nombre es siempre el actual en SQL, no una copia fosil en Mongo."
)

h2("5.3 Captura asincrona de eventos")

parrafo(
"El hook useTrackEvent envia los eventos al endpoint /api/cliente/eventos. "
"La insercion en Mongo es no-bloqueante respecto al renderizado: el "
"cliente nunca espera al log. Si MongoDB esta caido, el evento se pierde, "
"pero la experiencia web no se degrada. Esta decision refleja que el "
"log de eventos no es transaccionalmente critico."
)

h2("5.4 Generacion programada del feed")

parrafo(
"La regeneracion del feed se ejecuta on-demand a partir del TTL: cuando "
"el cliente entra y su feed caduco, la aplicacion dispara generarFeed. "
"La generacion en si misma cruza datos de Mongo (eventos, vistas) con "
"datos de SQL (productos comprados, stock por sede). El stock se consulta "
"contra Inventario.Stock_Actual a traves de los modulos data/ de la "
"aplicacion."
)

doc.add_page_break()

# =====================================================================
# 6. RESUMEN DE APORTES TECNICOS
# =====================================================================
h1("6. Resumen de aportes tecnicos del componente NoSQL")

parrafo(
"El componente NoSQL agrega valor en seis dimensiones distintas, cada una "
"asociada a una capacidad que MongoDB ofrece de forma nativa y que SQL "
"Server tendria que simular a un costo significativamente mayor."
)

tabla(
    ["Capacidad", "Aporte concreto al sistema"],
    [
        ["Validador $jsonSchema",
         "Integridad de resenas garantizada por el motor, sin duplicar codigo en aplicacion."],
        ["Subdocumentos embebidos",
         "Respuestas, fotos y datos opcionales viven dentro de la resena; sin tablas auxiliares ni JOINs."],
        ["Indice unico compuesto",
         "Una opinion por compra: la regla de negocio es una restriccion de motor, no codigo."],
        ["Coleccion time-series",
         "Escritura masiva de eventos con compresion columnar; agregaciones temporales eficientes."],
        ["TTL automatico",
         "Retencion de 90 dias en eventos y caducidad on-demand del feed, sin jobs de mantenimiento."],
        ["Vista materializada",
         "Lectura O(1) del feed personalizado, recalculado bajo demanda."],
    ]
)

h2("6.1 Cumplimiento de la rubrica")

parrafo(
"Conforme al criterio 'Modulo NoSQL (MongoDB u otro) — 50 puntos' de la "
"rubrica del proyecto, el componente descrito en este documento cubre los "
"cuatro requisitos de la valoracion Excelente:"
)

tabla(
    ["Requisito de la rubrica", "Cobertura en el sistema"],
    [
        ["Insercion de datos",
         "Scripts 01_datos_referencia_sql.js y 02_datos_mongo.js; insertarResena con verificacion contra SQL."],
        ["Consulta de datos",
         "Scripts 03_consultas_resenas.js, 04_consultas_eventos.js y 06_consultas_feed.js con 15+ consultas productivas."],
        ["Uso de estructuras adecuadas",
         "Documento con validador para resenas; time-series para eventos; vista materializada con TTL para el feed."],
        ["Consultas relevantes",
         "Cada consulta responde a una pregunta concreta del negocio: ranking, embudo, busquedas fallidas, recomendaciones."],
    ]
)

h2("6.2 Conclusion")

parrafo(
"El componente NoSQL del proyecto TiendaRopa no se incorporo por modernidad "
"sino por necesidad arquitectonica. Cada coleccion resuelve un problema "
"que el modelo relacional resolveria mal o a un costo desproporcionado. "
"Las decisiones tecnicas estan defendidas con razones de negocio: la regla "
"'una opinion por compra' como indice unico, la retencion de 90 dias como "
"TTL, la vigencia de 6 horas como contrato de frescura del feed. La "
"integracion con SQL Server respeta el rol de cada motor: el sistema "
"relacional sigue siendo la fuente de verdad transaccional, y MongoDB "
"actua como capa de enriquecimiento de la experiencia del cliente."
)

# =====================================================================
# Guardar
# =====================================================================
out = r"C:\AHS\Base de Datos 2\Tienda de Ropa\documentacion\Documentacion_NoSQL_MongoDB.docx"
doc.save(out)
print("OK:", out)
